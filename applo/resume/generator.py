import shutil
import httpx
from pathlib import Path
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from applo.utils.logger import logger


BULLET_CHARS = ("●", "•", "-", "*")


def _set_paragraph_text(para, new_text: str):
    """Replace paragraph text while preserving the first run's character formatting."""
    for run in para.runs:
        run.text = ""
    if para.runs:
        para.runs[0].text = new_text
    else:
        para.add_run(new_text)


def _patch_job_bullets(
    paragraphs: list,
    exp_indices: list[int],
    target_company: str,
    target_title: str,
    new_bullets: list[str],
) -> None:
    """Find the target job block in the experience section and replace its bullets."""
    # Find the paragraph index (within exp_indices) where the target job header appears
    job_start = None
    for pos, idx in enumerate(exp_indices):
        text = paragraphs[idx].text.lower()
        if target_company and target_company in text:
            job_start = pos
            break
        if target_title and target_title in text:
            job_start = pos
            break

    if job_start is None:
        return

    # Collect bullet paragraphs belonging to this job (until the next non-bullet, non-empty line)
    bullet_count = 0
    in_job = False
    for pos in range(job_start, len(exp_indices)):
        idx = exp_indices[pos]
        text = paragraphs[idx].text.strip()

        if not text:
            continue

        is_bullet = any(text.startswith(c) for c in BULLET_CHARS)

        if pos == job_start:
            in_job = True
            continue  # this is the header line, skip

        if in_job:
            if is_bullet:
                if bullet_count < len(new_bullets):
                    prefix = text[0]
                    _set_paragraph_text(paragraphs[idx], f"{prefix} {new_bullets[bullet_count]}")
                    bullet_count += 1
            else:
                # hit the next job header — stop
                break


def _to_pdf(docx_path: Path, pdf_path: Path) -> Path:
    from applo.config import settings
    url = f"{settings.gotenberg_url}/forms/libreoffice/convert"
    with open(docx_path, "rb") as f:
        response = httpx.post(
            url,
            files={"files": (docx_path.name, f, "application/vnd.openxmlformats-officedocument.wordprocessingml.document")},
            timeout=60,
        )
    if response.status_code != 200:
        raise RuntimeError(f"Gotenberg conversion failed: {response.status_code} {response.text}")
    pdf_path.write_bytes(response.content)
    return pdf_path


class ResumeGenerator:
    def generate(
        self,
        output_path: str | Path,
        master_docx_path: str | Path,
        resume_data: dict,
        optimized: dict,
    ) -> Path:
        """
        Copy master DOCX, patch optimized sections in-place, export to PDF.
        Layout and formatting are fully preserved from the original.
        """
        output_path = Path(output_path)
        output_path.parent.mkdir(parents=True, exist_ok=True)

        docx_out = output_path.with_suffix(".docx")
        shutil.copy2(master_docx_path, docx_out)

        doc = Document(docx_out)
        paragraphs = doc.paragraphs
        para_map: dict[str, list[int]] = resume_data.get("paragraph_map", {})

        # --- Summary ---
        summary_indices = para_map.get("professional summary") or para_map.get("summary", [])
        new_summary = optimized.get("summary", "")
        if summary_indices and new_summary:
            _set_paragraph_text(paragraphs[summary_indices[0]], new_summary)
            for idx in summary_indices[1:]:
                _set_paragraph_text(paragraphs[idx], "")

        # --- Skills ---
        skills_indices = para_map.get("skills") or para_map.get("technical skills", [])
        new_skills = optimized.get("skills", "")
        if skills_indices and new_skills:
            skill_lines = [l for l in new_skills.split("\n") if l.strip()]
            for i, idx in enumerate(skills_indices):
                _set_paragraph_text(paragraphs[idx], skill_lines[i] if i < len(skill_lines) else "")

        # --- Experience bullets ---
        exp_indices = (
            para_map.get("professional experience")
            or para_map.get("work experience")
            or para_map.get("experience", [])
        )
        bullet_data = optimized.get("experience_bullets", {})
        if isinstance(bullet_data, dict) and exp_indices:
            target_company = (bullet_data.get("company") or "").lower()
            target_title = (bullet_data.get("job_title") or "").lower()
            new_bullets = bullet_data.get("bullets", [])
            _patch_job_bullets(paragraphs, exp_indices, target_company, target_title, new_bullets)

        doc.save(docx_out)
        return _to_pdf(docx_out, output_path)

    def generate_cover_letter(
        self,
        output_path: str | Path,
        cover_letter_text: str,
        candidate_name: str,
    ) -> Path:
        output_path = Path(output_path)
        output_path.parent.mkdir(parents=True, exist_ok=True)
        docx_out = output_path.with_suffix(".docx")

        doc = Document()
        # Margins
        for section in doc.sections:
            section.top_margin = Pt(72)
            section.bottom_margin = Pt(72)
            section.left_margin = Pt(72)
            section.right_margin = Pt(72)

        name_para = doc.add_paragraph(candidate_name)
        name_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        name_run = name_para.runs[0]
        name_run.bold = True
        name_run.font.size = Pt(16)

        doc.add_paragraph()

        for block in cover_letter_text.split("\n\n"):
            if block.strip():
                p = doc.add_paragraph(block.strip())
                p.paragraph_format.space_after = Pt(10)

        doc.save(docx_out)
        return _to_pdf(docx_out, output_path)
